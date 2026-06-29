"""Cover letter renderer - layout only.

Reads LETTER_META and picked fragments from the application's selection.py,
merges with profile.py identity from the content package, and writes
james_kelley_cover_letter.docx into /work.

Podman invocation (engine from skill, content from resume repo):
  podman run --rm \\
    -v /path/to/application:/work \\
    -v /path/to/skill/scripts:/scripts \\
    -v /path/to/resume-repo/content:/content \\
    python:3.12-slim bash -c \\
    "pip install python-docx -q && python3 /scripts/build_cover_letter.py"

The script is called from /, so /scripts and /content are on sys.path automatically.
"""

import sys
import os

sys.path.insert(0, '/scripts')
sys.path.insert(0, '/')   # exposes /content as a top-level package

from docx import Document
from docx_helpers import (
    DARK, BLUE, GRAY,
    SIZE_NAME, SIZE_TITLE, SIZE_CONTACT, SIZE_BODY, SIZE_COMPANY,
    run, page_setup, section_header, add_bottom_border,
)

# ── Load content ──────────────────────────────────────────────────────────────
# /work contains the application's selection.py (LETTER_META + fragment picks)
sys.path.insert(0, '/work')

from content.profile import WHOS_RESUME_IS_THIS, CONTACT_LINE
import selection as sel

meta            = sel.LETTER_META       # {date, company, role, tone}
opening         = sel.OPENING           # one string (formatted with .format(**meta))
proofs          = sel.PROOFS            # list of strings
why             = sel.WHY_THIS_COMPANY  # free-text paragraph drafted per JD
closing         = sel.CLOSING           # one string
output_filename = getattr(sel, 'OUTPUT_FILENAME', 'james_kelley_cover_letter.docx')


def fmt(text):
    """Fill {company} and {role} placeholders from LETTER_META."""
    return text.format(company=meta['company'], role=meta['role'])


def body_para(doc, text, space_before=6, space_after=6):
    """Add a body-weight paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = __import__('docx').shared.Pt(space_before)
    p.paragraph_format.space_after  = __import__('docx').shared.Pt(space_after)
    run(p, text, bold=False, size=SIZE_BODY, color=DARK)
    return p


def build(work_dir='/work'):
    from docx.shared import Pt

    doc = Document()
    page_setup(doc)

    # ── Header: name ─────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run(p, WHOS_RESUME_IS_THIS, bold=True, size=SIZE_NAME, color=DARK)

    # ── Header: contact line ──────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run(p, CONTACT_LINE, size=SIZE_CONTACT, color=GRAY)
    add_bottom_border(p, color_hex='CCCCCC')

    # ── Date / top matter ─────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run(p, meta['date'], size=SIZE_BODY, color=GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run(p, meta['company'], bold=True, size=SIZE_BODY, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run(p, 're: ' + meta['role'], size=SIZE_BODY, color=BLUE)

    # ── Greeting ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run(p, 'Dear Hiring Team,', bold=False, size=SIZE_BODY, color=DARK)

    # ── Opening hook ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    run(p, fmt(opening), bold=False, size=SIZE_BODY, color=DARK)

    # ── Proof paragraphs ──────────────────────────────────────────────────────
    for proof in proofs:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(8)
        run(p, fmt(proof), bold=False, size=SIZE_BODY, color=DARK)

    # ── Why this company ──────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    run(p, fmt(why), bold=False, size=SIZE_BODY, color=DARK)

    # ── Closing paragraph ─────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(14)
    run(p, fmt(closing), bold=False, size=SIZE_BODY, color=DARK)

    # ── Sign-off ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run(p, 'Sincerely,', bold=False, size=SIZE_BODY, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run(p, WHOS_RESUME_IS_THIS.title(), bold=True, size=SIZE_BODY, color=DARK)

    # ── Document metadata ─────────────────────────────────────────────────────
    props = doc.core_properties
    props.author   = 'James Kelley'
    props.comments = 'Generated w/ ☕'

    # ── Output ────────────────────────────────────────────────────────────────
    out = os.path.join(work_dir, output_filename)
    doc.save(out)
    print('Saved: %s' % out)


if __name__ == '__main__':
    build()
